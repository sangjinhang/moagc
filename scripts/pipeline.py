#!/usr/bin/env python3
"""
COSMIC Excel → 完整工程需求文档（含增强内容 + 功能结构图 + 时序图 + H5 PRD重写 + 样式统一）
用法: python pipeline.py <输入.xlsx> [输出目录]
"""

import os, sys, re, json
import pandas as pd
from collections import OrderedDict

# ---- 命令行参数 ----
if len(sys.argv) < 2:
    print("用法: python pipeline.py <输入.xlsx> [输出目录]")
    sys.exit(1)

INPUT_XLSX = os.path.abspath(sys.argv[1])
OUTPUT_DIR = sys.argv[2] if len(sys.argv) > 2 else os.path.join(os.path.dirname(INPUT_XLSX), "output")
BASE_NAME = os.path.splitext(os.path.basename(INPUT_XLSX))[0]

# 本地脚本路径（云部署时所有脚本打包在同级 scripts/ 目录）
SCRIPTS_DIR = os.path.dirname(os.path.abspath(__file__))
SKILL_PATH = os.path.dirname(SCRIPTS_DIR)  # deploy/ 目录

os.makedirs(OUTPUT_DIR, exist_ok=True)

# 输出文件路径
WORD_PATH = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_需求描述.docx")
WORD_ENHANCED_PATH = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_需求描述.docx")  # 覆盖原文件
FSD_EXCEL_PATH = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_功能结构图.xlsx")
FSD_IMAGES_DIR = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_功能结构图_图片")
SEQ_DIR = os.path.join(OUTPUT_DIR, "时序图_H3_v3")
FINAL_WORD_PATH = os.path.join(OUTPUT_DIR, f"{BASE_NAME}_需求描述_完整版.docx")

print("=" * 60)
print("全流程: Excel → 增强需求文档 + 功能结构图 + 时序图")
print("=" * 60)

# ============================================================
# 步骤 0: 数据预处理（ffill 合并单元格）
# ============================================================
print("\n[步骤 0] 数据预处理...")
df = pd.read_excel(INPUT_XLSX)

# 规范化 OPEX 列名
for col in df.columns:
    if 'OPEX' in str(col) or '功能用户需求' in str(col):
        df.rename(columns={col: '功能用户需求'}, inplace=True)
        print(f"  列名映射: {repr(col)} → '功能用户需求'")
        break

merge_cols = ["一级模块", "二级模块", "三级模块"]
for col in merge_cols:
    df[col] = df[col].ffill()

# 10列格式适配：用功能过程填充四级模块（供后续时序图使用）
if "四级模块" not in df.columns and "功能过程" in df.columns:
    df["四级模块"] = df["功能过程"].ffill()

print(f"  数据行数: {len(df)}, 一级模块: {df['一级模块'].nunique()}, 二级模块: {df['二级模块'].nunique()}, 三级模块: {df['三级模块'].nunique()}")


# ============================================================
# 步骤 1: 生成纯文字 Word
# ============================================================
print("\n[步骤 1] 生成纯文字 Word 需求描述...")
sys.path.insert(0, SCRIPTS_DIR)
for mod_key in list(sys.modules.keys()):
    if "convert_cosmic13" in mod_key:
        del sys.modules[mod_key]

from convert_cosmic13 import convert_cosmic13_to_word
convert_cosmic13_to_word(INPUT_XLSX, WORD_PATH)
print(f"  完成: {WORD_PATH}")


# ============================================================
# 步骤 2: 增强 PRD 内容（H2/H3/H4 增强段落）
# ============================================================
print("\n[步骤 2] 增强 PRD 内容...")

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---- 构建数据索引 ----
l1_index = {}
for l1 in df['一级模块'].dropna().unique():
    l1_rows = df[df['一级模块'] == l1]
    l1_index[l1] = {
        'l2_names': [x for x in l1_rows['二级模块'].dropna().unique().tolist() if str(x).strip() and str(x).strip() != 'nan'],
        'l3_count': l1_rows['三级模块'].dropna().nunique(),
    }

h3_to_l1 = {}
current_l1 = ''
for _, row in df.drop_duplicates(subset=['二级模块'], keep='first').iterrows():
    if pd.notna(row.get('一级模块')):
        current_l1 = row['一级模块']
    if pd.notna(row.get('二级模块')):
        h3_to_l1[row['二级模块']] = current_l1

l2_index = {}
for l2 in df['二级模块'].dropna().unique():
    if not str(l2).strip() or str(l2).strip() == 'nan':
        continue
    l2_rows = df[df['二级模块'] == l2]
    l2_index[l2] = {
        'l1_name': h3_to_l1.get(l2, ''),
        'l3_count': l2_rows['三级模块'].dropna().nunique(),
        'l3_names': [x for x in l2_rows['三级模块'].dropna().unique().tolist() if str(x).strip() and str(x).strip() != 'nan'],
    }

l3_index = {}
for _, row in df.drop_duplicates(subset=['一级模块','二级模块','三级模块'], keep='first').iterrows():
    key = (row['一级模块'], row['二级模块'], row['三级模块'])
    l3_index[key] = {
        '功能用户': str(row.get('功能用户', '')).replace('\n', '；') if pd.notna(row.get('功能用户')) else '',
        '功能用户需求': str(row.get('功能用户需求', '')) if pd.notna(row.get('功能用户需求')) else '',
        '触发事件': str(row.get('触发事件', '')) if pd.notna(row.get('触发事件')) else '',
        '功能过程': str(row.get('功能过程', '')) if pd.notna(row.get('功能过程')) else '',
    }

# ---- 收集二级/三级模块业务字段 ----
l2_biz = {}
for l2 in df['二级模块'].dropna().unique():
    if not str(l2).strip() or str(l2).strip() == 'nan':
        continue
    rows = df[df['二级模块'] == l2]
    l2_biz[l2] = {
        'triggers': [r for r in rows['触发事件'].dropna().unique().tolist() if r and str(r).strip() and str(r).strip() != 'nan'],
        'processes': [r for r in rows['功能过程'].dropna().unique().tolist() if r and str(r).strip() and str(r).strip() != 'nan'],
        'data_groups': [r for r in rows['数据组'].dropna().unique().tolist() if r and str(r).strip() and str(r).strip() != 'nan'],
    }

l3_biz = {}
for key, info in l3_index.items():
    l1, l2, l3 = key
    rows = df[(df['一级模块']==l1)&(df['二级模块']==l2)&(df['三级模块']==l3)]
    l3_biz[key] = {
        'data_groups': [r for r in rows['数据组'].dropna().unique().tolist() if r and str(r).strip() and str(r).strip() != 'nan'],
        'processes': [r for r in rows['功能过程'].dropna().unique().tolist() if r and str(r).strip() and str(r).strip() != 'nan'],
    }

# ---- 角色提取函数 ----
def extract_roles(raw_text):
    roles = set()
    for part in raw_text.replace('\n', '；').split('；'):
        part = part.strip()
        for prefix in ['发起者：', '接收者：', '发起者:', '接收者:']:
            if part.startswith(prefix):
                role = part[len(prefix):].strip()
                if role and role not in ['服务端', '系统']:
                    roles.add(role)
    return list(roles)

# ---- H2 业务目标预设 ----
goal_map = {
    'MobileClaw': '为用户提供稳定、高效的OpenClaw会话接入与消息交互体验，支撑AI助手在日常办公场景中的实际落地。',
    'AI': '提升群聊场景下的智能化服务水平，让用户在群组内即可完成智能问答、文件处理等操作，减少跨系统切换。',
    '全双工': '构建全双工实时通信架构，为弹窗类业务提供底层通信支撑，确保多端协同场景下的消息实时性与一致性。',
    '智信能力提升': '强化智信在聊天、文件、安全等方面的基础能力，为上层业务提供更可靠、更灵活的通讯底座。',
    '产品能力提升': '完善服务号、外部用户等产品能力，丰富平台面向不同用户群体的服务场景与功能覆盖。',
    '分级运营门户': '搭建分级运营门户，满足一二级平台不同层级管理人员对运营数据、权限配置、安全管控的差异化需求。',
    '二级门户数据运营': '为二级运营门户提供指标分析、报表展示等数据运营能力，支撑省专层面的精细化运营决策。',
    '智信运营能力提升': '增强智信在群组管理、消息管控、快通知审核等方面的运营管理能力，满足企业级通讯安全合规要求。',
    '智能运营': '构建智能客服、灵犀助手、机器人群聊等智能运营能力，通过AI技术提升运营效率与服务质量。',
    '信网项目': '推进信网在鸿蒙、PC等端的客户端建设，完成管理平台分权分域能力升级，拓展信网产品的终端覆盖范围。',
}

# ---- 增强内容生成函数 ----
def gen_l1(l1_name, info):
    names = info['l2_names']
    n = len(names)
    overview = f"【模块概述】{l1_name}作为本期工程建设的重点模块，下设{n}个二级子模块，覆盖{info['l3_count']}项功能需求。"
    if n >= 3:
        overview += f"主要围绕{names[0]}、{names[1]}、{names[2]}等领域展开。"
    elif n == 2:
        overview += f"主要围绕{names[0]}和{names[1]}两个领域展开。"
    else:
        overview += f"主要围绕{names[0]}领域展开。"

    goal_text = None
    for k, v in goal_map.items():
        if k in l1_name:
            goal_text = v; break
    if not goal_text:
        goal_text = f"通过{n}个子模块的协同建设，为平台用户提供更加完善的业务支撑能力，推动产品体验与运营效率的持续提升。"

    lines = [overview, f"【业务目标】{goal_text}"]
    lines.append(f"【功能范围】{l1_name}包含以下{n}个二级模块：")
    for i, nm in enumerate(names):
        lines.append(f"（{i+1}）{nm}")

    has_mobile = any(any(k in nm for k in ['Android','iOS','鸿蒙','PC','移动端','客户端适配']) for nm in names)
    has_config = any(any(k in nm for k in ['配置','管理','规则']) for nm in names)
    has_ai = any(any(k in nm for k in ['AI','智能','机器人','问答']) for nm in names)
    points = []
    if has_ai:
        points.append("AI能力的准确性、响应速度与用户体验需要持续优化，特别是在文件处理、问答质量等核心场景上")
    if has_mobile:
        points.append("多端适配方面需确保Android、iOS、鸿蒙、PC各端的功能一致性与体验流畅度")
    if has_config:
        points.append("配置管理类功能需要兼顾灵活性与易用性，让运营人员能够高效完成规则配置与生效管理")
    if not points:
        points.append("各子模块之间需要保持数据流转与状态同步的一致性，避免因边界场景导致功能异常")
        points.append("需充分考虑高并发、大数据量场景下的系统性能表现，确保核心功能的稳定可用")
    lines.append(f"【建设要点】{'；'.join(points)}。")
    return '\n'.join(lines)


def gen_l2(l2_name, info):
    l1 = info['l1_name']
    l3_count = info['l3_count']
    l3_names = info['l3_names']
    biz = l2_biz.get(l2_name, {})
    lines = []

    pos = f"【模块定位】{l2_name}是{l1}下的"
    if l3_count <= 3:
        pos += f"基础功能单元，包含{l3_count}项功能需求。"
    elif l3_count <= 8:
        pos += f"核心功能单元，包含{l3_count}项功能需求。"
    else:
        pos += f"重要功能单元，共涵盖{l3_count}项功能需求。"

    # 按模块名称语义推断业务定位
    if any(k in l2_name for k in ['配置','管理','规则']):
        pos += "主要为运营管理人员提供后台配置与管控能力。"
    elif any(k in l2_name for k in ['AI','智能','机器人','问答']):
        pos += "主要面向用户提供智能化、自动化的业务处理与交互体验。"
    elif any(k in l2_name for k in ['客户端适配','Android','iOS','鸿蒙移动端','MacPC','WinPC','国产化PC']):
        pos += "主要负责对应终端平台的功能适配与体验优化，确保管理端配置在各客户端上的正确生效。"
    elif any(k in l2_name for k in ['指标','报表','数据','分析']):
        pos += "主要支撑运营数据的采集、分析与可视化展示，为运营决策提供数据依据。"
    elif any(k in l2_name for k in ['接入','会话','消息','链路','语音']):
        pos += "主要负责用户会话的建立、消息的收发以及交互过程的管控。"
    elif any(k in l2_name for k in ['安全','权限','审计','日志']):
        pos += "主要负责平台安全管控、权限治理与操作审计。"
    elif any(k in l2_name for k in ['服务号']):
        pos += "主要负责服务号的运营管理，包括内容发布、消息推送、用户订阅等全生命周期管理。"
    elif any(k in l2_name for k in ['门户','通讯录','VIP','基础能力']):
        pos += "主要提供平台基础支撑能力，确保各业务模块的正常运行。"
    else:
        pos += "面向用户提供该业务场景下的完整功能支撑。"
    lines.append(pos)

    # 用户角色
    actors = set()
    for (l1_, l2_, l3_), inf in l3_index.items():
        if l2_ == l2_name and inf['功能用户']:
            roles = extract_roles(inf['功能用户'])
            for r in roles: actors.add(r)
    if actors:
        actor_list = sorted(actors)
        role_text = '、'.join(actor_list[:4])
        has_user = any('用户' in a for a in actor_list)
        has_admin = any(any(k in a for k in ['管理','运维','运营','操作']) for a in actor_list)
        if has_user and has_admin:
            desc = f"面向{role_text}等角色，分别提供前台业务操作与后台管理维护能力。"
        elif has_admin:
            desc = f"面向{role_text}等角色，主要提供后台管理与运营维护能力。"
        else:
            desc = f"面向{role_text}等角色提供服务。"
        lines.append(f"【用户角色】{desc}")

    lines.append(f"【功能构成】{l2_name}由以下{l3_count}个功能单元组成：")
    for i, nm in enumerate(l3_names):
        lines.append(f"（{i+1}）{nm}")

    # 核心流程
    triggers = biz.get('triggers', [])
    clean_triggers = []
    for t in triggers[:8]:
        t = str(t).strip()
        if t and t != 'nan':
            t_clean = re.sub(r'^(用户|管理人员|运营人员)\s*(提交|发起)?\s*', '', t).strip()
            if t_clean and t_clean not in clean_triggers and len(clean_triggers) < 4:
                clean_triggers.append(t_clean)

    flow_text = f"【核心流程】{l2_name}的主要业务场景包括："
    if clean_triggers:
        flow_text += '、'.join(clean_triggers) + '。'
    else:
        processes = biz.get('processes', [])
        clean_procs = [str(p).strip() for p in processes[:4] if str(p).strip() and str(p).strip() != 'nan']
        if clean_procs:
            flow_text += '、'.join(clean_procs) + '。'
        else:
            flow_text += f"覆盖{l3_names[0]}至{l3_names[-1]}共{l3_count}项功能，用户按需触发相应操作即可。"

    if any(k in l2_name for k in ['配置','管理','规则']):
        flow_text += "管理人员在后台完成配置后，规则实时下发生效。"
    elif any(k in l2_name for k in ['查询','查看','检索','报表']):
        flow_text += "系统校验请求方权限后查询数据并整理返回。"
    elif any(k in l2_name for k in ['适配','客户端','Android','iOS','鸿蒙']):
        flow_text += "管理端配置变更后，终端通过版本更新或配置同步机制获取最新规则并生效。"
    elif any(k in l2_name for k in ['群聊','群组','机器人']):
        flow_text += "群内成员通过@机器人或直接发送消息触发交互，系统识别意图后执行对应处理。"
    lines.append(flow_text)
    return '\n'.join(lines)


def gen_l3(l1_name, l2_name, l3_name, info):
    lines = []
    if info['功能用户']:
        roles = extract_roles(info['功能用户'])
        if roles:
            lines.append(f"【用户角色】面向{'、'.join(roles)}提供服务。")
        else:
            lines.append(f"【用户角色】{info['功能用户']}。")
    if info['功能用户需求'] and info['功能用户需求'] != 'nan':
        lines.append(f"【业务需求】{info['功能用户需求']}。")
    if info['触发事件'] and info['触发事件'] != 'nan':
        trigger = re.sub(r'^(用户|管理人员|运营人员)\s*(提交|发起)?\s*', '', str(info['触发事件']).strip())
        lines.append(f"【触发条件】{trigger}。")
    if info['功能过程'] and info['功能过程'] != 'nan':
        lines.append(f"【处理流程】{info['功能过程']}。")
    key = (l1_name, l2_name, l3_name)
    biz = l3_biz.get(key, {})
    dgs = [d for d in biz.get('data_groups', []) if d and str(d).strip() and str(d).strip() != 'nan']
    if dgs and len(dgs) >= 2:
        lines.append(f"【涉及数据】主要涉及{'、'.join(dgs[:4])}等业务数据。")
    elif len(dgs) == 1:
        lines.append(f"【涉及数据】主要涉及{dgs[0]}。")
    return '\n'.join(lines)


# ---- 插入增强段落 ----
def add_enh(elem, text):
    """在 elem 后插入增强段落"""
    for part in reversed(text.split('\n')):
        if not part.strip():
            continue
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        if part.startswith('【') and '】' in part:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLineChars'), '0')
            ind.set(qn('w:firstLine'), '0')
        else:
            ind = OxmlElement('w:ind')
            ind.set(qn('w:firstLineChars'), '200')
            ind.set(qn('w:firstLine'), '420')
        pPr.append(ind)
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'), '60')
        sp.set(qn('w:after'), '60')
        pPr.append(sp)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        for attr in [qn('w:ascii'), qn('w:eastAsia'), qn('w:hAnsi')]:
            rFonts.set(attr, "宋体")
        rPr.append(rFonts)
        for tag in ['w:sz', 'w:szCs']:
            el = OxmlElement(tag)
            el.set(qn('w:val'), '21')
            rPr.append(el)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
        t.text = part
        r.append(t)
        p.append(r)
        elem.addnext(p)


doc = Document(WORD_PATH)

# 收集标题 → 需求说明段落
entries = []
paras = doc.paragraphs
for i, p in enumerate(paras):
    if p.style.name in ('Heading 2', 'Heading 3', 'Heading 4'):
        title = p.text.strip()
        if not title or title == 'nan':
            continue
        for j in range(i + 1, min(i + 5, len(paras))):
            if paras[j].style.name == 'Normal' and '需求说明' in paras[j].text:
                entries.append({'di': j, 's': p.style.name, 't': title})
                break

# 从后往前插入
counts = {'H2': 0, 'H3': 0, 'H4': 0}
for idx in range(len(entries) - 1, -1, -1):
    e = entries[idx]
    paras = doc.paragraphs
    if e['di'] >= len(paras):
        continue
    desc_elem = paras[e['di']]._element
    enh = None
    if e['s'] == 'Heading 2' and e['t'] in l1_index:
        enh = gen_l1(e['t'], l1_index[e['t']])
        counts['H2'] += 1
    elif e['s'] == 'Heading 3' and e['t'] in l2_index:
        enh = gen_l2(e['t'], l2_index[e['t']])
        counts['H3'] += 1
    elif e['s'] == 'Heading 4':
        for (l1, l2, l3), inf in l3_index.items():
            if l3 == e['t']:
                enh = gen_l3(l1, l2, e['t'], inf)
                counts['H4'] += 1
                break
    if enh:
        add_enh(desc_elem, enh)

# 清理COSMIC术语
for p in doc.paragraphs:
    for run in p.runs:
        t = run.text
        if 'COSMIC功能点估算表' in t:
            run.text = t.replace('COSMIC功能点估算表', '功能拆分表')
        if '每项功能点均包含' in t:
            run.text = t.replace('每项功能点均包含', '每项功能均包含')

doc.save(WORD_ENHANCED_PATH)
print(f"  增强完成: H2={counts['H2']}, H3={counts['H3']}, H4={counts['H4']}")


# ============================================================
# 步骤 3: 生成功能结构图 Excel
# ============================================================
print("\n[步骤 3] 生成功能结构图 Excel...")
l1_to_l2 = OrderedDict()
l2_to_l3 = OrderedDict()

for _, row in df.iterrows():
    l1, l2, l3 = row['一级模块'], row['二级模块'], row['三级模块']
    if pd.isna(l1) or pd.isna(l2) or pd.isna(l3):
        continue
    l1, l2, l3 = str(l1).strip(), str(l2).strip(), str(l3).strip()
    if not l1 or not l2 or not l3 or l1 == 'nan' or l2 == 'nan' or l3 == 'nan':
        continue
    if l1 not in l1_to_l2:
        l1_to_l2[l1] = []
    if l2 not in l1_to_l2[l1]:
        l1_to_l2[l1].append(l2)
    if l2 not in l2_to_l3:
        l2_to_l3[l2] = []
    if l3 not in l2_to_l3[l2]:
        l2_to_l3[l2].append(l3)

sys.path.insert(0, SCRIPTS_DIR)
for mod_key in list(sys.modules.keys()):
    if "generate_fsd_excel" in mod_key:
        del sys.modules[mod_key]

from generate_fsd_excel import generate_fsd_excel_from_mapping
fsd_result = generate_fsd_excel_from_mapping(l1_to_l2, l2_to_l3, FSD_EXCEL_PATH)
print(f"  完成: {FSD_EXCEL_PATH} ({fsd_result['l1_count']}个一级, {fsd_result['l2_count']}个二级)")


# ============================================================
# 步骤 4: 功能结构图 Excel → PNG
# ============================================================
print("\n[步骤 4] 功能结构图 → PNG 图片...")
for mod_key in list(sys.modules.keys()):
    if "convert_fsd_to_images" in mod_key:
        del sys.modules[mod_key]

from convert_fsd_to_images import convert_fsd_to_images
img_result = convert_fsd_to_images(FSD_EXCEL_PATH, FSD_IMAGES_DIR)
print(f"  完成: {FSD_IMAGES_DIR} ({len(img_result['images'])}张)")


# ============================================================
# 步骤 5: 批量生成时序图 + 描述
# ============================================================
print("\n[步骤 5] 生成时序图 + 描述...")
os.makedirs(SEQ_DIR, exist_ok=True)

sys.path.insert(0, SCRIPTS_DIR)
for mod_key in list(sys.modules.keys()):
    if "draw_sequence" in mod_key:
        del sys.modules[mod_key]

from draw_sequence import derive_participants, derive_messages, generate_description, generate_seq_diagram

descriptions = {}
l2_list = [x for x in df['二级模块'].dropna().unique() if str(x).strip() and str(x).strip() != 'nan']
total_l2 = len(l2_list)

for idx, l2 in enumerate(l2_list):
    rows = df[df['二级模块'] == l2]
    participants = derive_participants(l2, rows)
    messages = derive_messages(l2, rows, participants)
    desc = generate_description(l2, rows, participants, messages)
    descriptions[l2] = desc

    safe_name = l2.replace('/', '_').replace('\\', '_').replace('?', '').replace(':', '')
    generate_seq_diagram(l2, participants, messages, os.path.join(SEQ_DIR, f"{safe_name}_时序图.png"))
    if (idx + 1) % 5 == 0:
        print(f"  进度: {idx+1}/{total_l2} ({l2})")

print(f"  完成: {total_l2}个二级模块的时序图")

with open(os.path.join(SEQ_DIR, "descriptions.json"), 'w', encoding='utf-8') as f:
    json.dump(descriptions, f, ensure_ascii=False, indent=2)


# ============================================================
# 步骤 6: 组装 Word（功能结构图 + 时序图 + 描述）
# ============================================================
print("\n[步骤 6] 组装最终 Word 文档...")

from docx.shared import Cm
from PIL import Image
from docx.text.paragraph import Paragraph

doc = Document(WORD_ENHANCED_PATH)

with open(os.path.join(SEQ_DIR, "descriptions.json"), 'r', encoding='utf-8') as f:
    descriptions = json.load(f)

fsd_files = {}
if os.path.exists(FSD_IMAGES_DIR):
    for f in os.listdir(FSD_IMAGES_DIR):
        if f.endswith('.png'):
            fsd_files[f.replace('.png', '')] = os.path.join(FSD_IMAGES_DIR, f)

seq_files = {}
if os.path.exists(SEQ_DIR):
    for f in os.listdir(SEQ_DIR):
        if f.endswith('_时序图.png'):
            seq_files[f.replace('_时序图.png', '')] = os.path.join(SEQ_DIR, f)

def make_text_p(text, bold=False, size=10.5):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    for attr in [qn('w:ascii'), qn('w:eastAsia'), qn('w:hAnsi')]:
        rFonts.set(attr, "宋体")
    rPr.append(rFonts)
    for tag in ['w:sz', 'w:szCs']:
        el = OxmlElement(tag)
        el.set(qn('w:val'), str(int(size * 2)))
        rPr.append(el)
    if bold:
        rPr.append(OxmlElement('w:b'))
    r.append(rPr)
    t = OxmlElement('w:t')
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    t.text = text
    r.append(t)
    p.append(r)
    return p

def make_img_p(img_path, max_cm=14):
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc')
    jc.set(qn('w:val'), 'center')
    pPr.append(jc)
    p.append(pPr)
    with Image.open(img_path) as im:
        wpx, hpx = im.size
    wcm = min(wpx * 0.0265, max_cm)
    tmp = Paragraph(p, doc)
    run = tmp.add_run()
    run.add_picture(img_path, width=Cm(wcm))
    return p

# H2: 功能结构图 addnext
paras = doc.paragraphs
h2_entries = [(i, p.text.strip()) for i, p in enumerate(paras) if p.style.name == 'Heading 2' and p.text.strip() and p.text.strip() != 'nan']
for idx in range(len(h2_entries) - 1, -1, -1):
    i, name = h2_entries[idx]
    fsd = fsd_files.get(name)
    if fsd and os.path.exists(fsd):
        try:
            paras[i]._element.addnext(make_img_p(fsd))
        except Exception:
            pass

# H3: 功能结构图 + 时序图 + 描述
paras = doc.paragraphs
h3_entries = [(i, p.text.strip()) for i, p in enumerate(paras) if p.style.name == 'Heading 3' and p.text.strip() and p.text.strip() != 'nan']

for idx in range(len(h3_entries) - 1, -1, -1):
    i, l2_name = h3_entries[idx]
    fsd = fsd_files.get(l2_name)
    seq = seq_files.get(l2_name)
    desc = descriptions.get(l2_name, [])

    need_desc_elem = None
    for j in range(i + 1, min(i + 12, len(paras))):
        if paras[j].style.name == 'Normal' and '需求说明' in paras[j].text:
            need_desc_elem = paras[j]._element
            break
    if need_desc_elem is None:
        continue

    # 功能结构图插在需求说明前面
    if fsd and os.path.exists(fsd):
        try:
            need_desc_elem.addprevious(make_img_p(fsd))
        except Exception:
            pass

    # 时序图 + 描述插在增强内容末尾
    enhancement_elem = need_desc_elem.getnext()
    if enhancement_elem is not None:
        elem_text = ''.join(r.text or '' for r in enhancement_elem.findall('.//' + qn('w:t')))
        if '【' in elem_text and '】' in elem_text:
            last_elem = enhancement_elem
            while True:
                next_elem = last_elem.getnext()
                if next_elem is None:
                    break
                next_text = ''.join(r.text or '' for r in next_elem.findall('.//' + qn('w:t')))
                next_style = next_elem.find('.//' + qn('w:pStyle'))
                if next_style is not None:
                    val = next_style.get(qn('w:val'))
                    if val and val.startswith('Heading'):
                        break
                if next_text.strip().startswith('需求说明'):
                    break
                last_elem = next_elem
            if seq and os.path.exists(seq):
                try:
                    seq_elem = make_img_p(seq)
                    last_elem.addnext(seq_elem)
                    last_elem = seq_elem
                except Exception:
                    pass
            if desc:
                try:
                    title_elem = make_text_p("时序图描述：", bold=True)
                    last_elem.addnext(title_elem)
                    last_elem = title_elem
                    for d_idx in range(len(desc)):
                        item_elem = make_text_p(f"{d_idx + 1}. {desc[d_idx]}")
                        last_elem.addnext(item_elem)
                        last_elem = item_elem
                except Exception:
                    pass

# 行距统一 1.3 倍
for p in doc.paragraphs:
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)
    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:line'), '312')
    spacing.set(qn('w:lineRule'), 'auto')

doc.save(FINAL_WORD_PATH)
print(f"  完成: {FINAL_WORD_PATH}")


# ============================================================
# 步骤 7: H5 功能过程重写为 PRD 风格
# ============================================================
print("\n[步骤 7] H5 功能过程重写为 PRD 风格...")
# 使用绝对路径（expanduser 在沙箱中返回沙箱路径）
REAL_USER_SKILLS = SCRIPTS_DIR
sys.path.insert(0, SCRIPTS_DIR)
for mod_key in list(sys.modules.keys()):
    if "rewrite_h5" in mod_key:
        del sys.modules[mod_key]

try:
    from rewrite_h5_to_prd import rewrite_h5_to_prd
    rewrite_h5_to_prd(input_xlsx=INPUT_XLSX, input_docx=FINAL_WORD_PATH, output_docx=FINAL_WORD_PATH)
    print(f"  H5重写完成")
except ImportError as e:
    print(f"  [跳过] rewrite_h5_to_prd 模块未找到: {e}，已有Word保持原样")

# ============================================================
# 步骤 8: 统一样式（字体宋体10.5pt、缩进2字符、段前段后60、行距1.3）
# ============================================================
print("\n[步骤 8] 统一样式...")
from docx.oxml import OxmlElement

doc = Document(FINAL_WORD_PATH)
FONT_FAMILY = "宋体"
FONT_SIZE = "21"
INDENT_LINE = "420"
INDENT_CHARS = "200"

normalized = 0
for p in doc.paragraphs:
    if '需求说明' not in p.text:
        continue
    if p.text.strip().startswith('【'):
        continue

    pPr = p._element.find(qn('w:pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        p._element.insert(0, pPr)

    ind = pPr.find(qn('w:ind'))
    if ind is None:
        ind = OxmlElement('w:ind')
        pPr.append(ind)
    for attr_name in ['firstLine', 'firstLineChars', 'left', 'leftChars', 'right', 'rightChars', 'hanging', 'hangingChars']:
        if ind.get(qn('w:' + attr_name)) is not None:
            del ind.attrib[qn('w:' + attr_name)]
    ind.set(qn('w:firstLineChars'), INDENT_CHARS)
    ind.set(qn('w:firstLine'), INDENT_LINE)

    spacing = pPr.find(qn('w:spacing'))
    if spacing is None:
        spacing = OxmlElement('w:spacing')
        pPr.append(spacing)
    spacing.set(qn('w:before'), '60')
    spacing.set(qn('w:after'), '60')
    spacing.set(qn('w:line'), '312')
    spacing.set(qn('w:lineRule'), 'auto')

    for run in p.runs:
        rPr = run._element.find(qn('w:rPr'))
        if rPr is None:
            rPr = OxmlElement('w:rPr')
            run._element.insert(0, rPr)
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = OxmlElement('w:rFonts')
            rPr.append(rFonts)
        for attr in ['ascii', 'eastAsia', 'hAnsi', 'cs']:
            rFonts.set(qn('w:' + attr), FONT_FAMILY)
        for tag in ['w:sz', 'w:szCs']:
            sz_el = rPr.find(qn(tag))
            if sz_el is None:
                sz_el = OxmlElement(tag)
                rPr.append(sz_el)
            sz_el.set(qn('w:val'), FONT_SIZE)
        color = rPr.find(qn('w:color'))
        if color is not None:
            rPr.remove(color)
    normalized += 1

doc.save(FINAL_WORD_PATH)
print(f"  规范化完成: {normalized} 个需求说明段落已统一样式")

# ============================================================
# 完成
# ============================================================
print(f"\n{'=' * 60}")
print(f"全流程完成！")
print(f"  最终文档: {FINAL_WORD_PATH}")
file_size_mb = os.path.getsize(FINAL_WORD_PATH) / 1024 / 1024 if os.path.exists(FINAL_WORD_PATH) else 0
print(f"  文件大小: {file_size_mb:.1f} MB")
print(f"{'=' * 60}")
