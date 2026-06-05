#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
批量重写功能过程(H5)内容为PM风格PRD：
需求说明 + 叙述段落 + 流程数据（带序号①②③④的分项明细）
"""

import pandas as pd
import re
import os
from collections import defaultdict
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def rewrite_h5_to_prd(input_xlsx, input_docx, output_docx):
    """
    Main entry point: rewrite all H5 sections in a Word document
    using PM-style narrative + breakdown format.

    Args:
        input_xlsx: Source COSMIC Excel file (10-column format)
        input_docx: Assembled Word document with H5 from convert_cosmic13
        output_docx: Output path for rewritten document
    """
    # ============================================================
    # Phase 1: Extract function process data from Excel
    # ============================================================
    print("[H5 Rewrite] Loading Excel...")
    df = pd.read_excel(input_xlsx)
    for col in ['一级模块', '二级模块', '三级模块', '功能过程']:
        df[col] = df[col].ffill()

    fp_data = defaultdict(lambda: {'l1': '', 'l2': '', 'l3': '', 'trigger': '', 'user': '', 'subs': []})
    for _, row in df.iterrows():
        fp_name = str(row['功能过程']).strip()
        if not fp_name or fp_name == 'nan':
            continue
        entry = fp_data[fp_name]
        entry['l1'] = str(row.get('一级模块', '')) if pd.notna(row.get('一级模块')) else ''
        entry['l2'] = str(row.get('二级模块', '')) if pd.notna(row.get('二级模块')) else ''
        entry['l3'] = str(row.get('三级模块', '')) if pd.notna(row.get('三级模块')) else ''
        if pd.notna(row.get('触发事件')):
            entry['trigger'] = str(row['触发事件'])
        if pd.notna(row.get('功能用户')):
            entry['user'] = str(row['功能用户'])
        sub_desc = str(row.get('子过程描述', '')).strip()
        if sub_desc and sub_desc != 'nan':
            dg = str(row.get('数据组', '')).strip() if pd.notna(row.get('数据组')) else ''
            da = str(row.get('数据属性', '')).strip() if pd.notna(row.get('数据属性')) else ''
            entry['subs'].append({
                'desc': sub_desc,
                'dg': dg if dg != 'nan' else '',
                'da': da if da != 'nan' else ''
            })
    print(f"[H5 Rewrite] Extracted {len(fp_data)} function processes")

    # ============================================================
    # Phase 2: Build PRD paragraphs
    # ============================================================

    def clean_trigger(trigger):
        return re.sub(
            r'^(用户|运营人员|管理人员|管理员|操作员|开发者|群成员)\s*(提交|发起|请求|点击|输入|选择|配置|设定)?\s*',
            '', trigger.strip()
        )

    def is_input_sub(desc):
        return bool(re.match(r'^(输入|提交|录入|上传|填写|发起|选择|配置|设定|设置|创建)', desc))

    def gen_paragraphs(fp_name, entry):
        """
        New format:
        1. 需求说明：one-sentence summary
        2. ①②③... each sub-process as a full PM-style paragraph
        No bold, no 流程数据 heading, all plain text.
        """
        trigger = entry.get('trigger', '')
        subs = entry['subs']
        scene = clean_trigger(trigger) if trigger and trigger != 'nan' else fp_name

        if not subs:
            return [f"需求说明：{scene}时，系统自动完成相应处理。"]

        # Extract actor
        user_text = entry.get('user', '')
        actor = ''
        if user_text and user_text != 'nan':
            parts = user_text.split('\n')
            for p in parts:
                for prefix in ['发起者：', '发起者:', '接收者：', '接收者:']:
                    if p.strip().startswith(prefix):
                        actor = p.strip()[len(prefix):].strip()
                        break
                if actor:
                    break
        if not actor:
            actor = '用户'

        num = ["①", "②", "③", "④", "⑤", "⑥", "⑦", "⑧", "⑨", "⑩", "⑪", "⑫", "⑬", "⑭", "⑮", "⑯", "⑰", "⑱", "⑲", "⑳"]

        paragraphs = []

        # ==== Summary line: PM-style narrative, NO statistics ====
        dg_names = []
        for s in subs:
            if s['dg'] and s['dg'] not in dg_names:
                dg_names.append(s['dg'])
        dg_summary = dg_names[0] if len(dg_names) == 1 else '、'.join(dg_names[:3])

        # Build narrative summary based on sub-process flow
        input_descs = [s['desc'] for s in subs if is_input_sub(s['desc'])]
        read_descs = [s['desc'] for s in subs if re.match(r'^(读取|查询|检索|获取|加载)', s['desc'])]
        output_descs = [s['desc'] for s in subs if re.match(r'^(返回|展示|输出|生成|推送|通知)', s['desc'])]
        process_descs = [s['desc'] for s in subs if not is_input_sub(s['desc']) and not re.match(r'^(读取|查询|检索|获取|加载|返回|展示|输出|生成|推送|通知)', s['desc'])]

        # Build a PM-style one-sentence narrative (scene-driven, no statistics)
        if input_descs and read_descs and output_descs:
            summary = f"当{actor}发起{scene}的请求时，服务端接收相关请求参数，对系统已维护的{dg_summary}进行检索与筛选，完成数据匹配后提取指定的业务字段，最终将整理完成的处理结果返回至客户端。"
        elif input_descs and output_descs:
            summary = f"当{actor}发起{scene}的请求时，服务端接收请求参数后进行相应处理，并将处理结果返回至客户端。"
        elif read_descs and output_descs:
            summary = f"当{actor}发起{scene}的请求时，服务端检索{dg_summary}并返回匹配结果。"
        elif process_descs:
            summary = f"当{actor}发起{scene}的请求时，服务端按照业务规则完成相应处理。"
        else:
            summary = f"当{actor}发起{scene}的请求时，系统按既定流程完成处理。"

        paragraphs.append(f"需求说明：{summary}")

        # ==== Sub-process paragraphs: each as a complete PM-style sentence ====
        for i, s in enumerate(subs):
            desc = s['desc']
            attrs = [a.strip() for a in s['da'].split(',') if a.strip()] if s['da'] else []
            dg = s['dg']
            attr_text = '、'.join(attrs) if attrs else ''

            if is_input_sub(desc):
                if dg and attr_text:
                    sub_para = f"当{actor}发起{scene}时，服务端接收{dg}，解析请求中携带的{attr_text}。"
                elif dg:
                    sub_para = f"当{actor}发起{scene}时，服务端接收{dg}。"
                else:
                    sub_para = f"当{actor}发起{scene}时，服务端接收{desc}。"
            elif re.match(r'^(读取|查询|检索|获取|加载)', desc):
                if dg and attr_text:
                    sub_para = f"随后依托接收到的查询条件，服务端对系统内已维护的{dg}进行精准检索，从匹配结果中提取{attr_text}。"
                elif dg:
                    sub_para = f"随后依托接收到的查询条件，服务端对系统内已维护的{dg}进行精准检索。"
                else:
                    sub_para = f"随后服务端{desc}。"
            elif re.match(r'^(返回|展示|输出|生成|推送|通知)', desc):
                if dg and attr_text:
                    sub_para = f"完成数据提取后，服务端按照统一格式完成数据组装封装，将{dg}（{attr_text}）原路返回至客户端，供前端进行页面渲染与内容展示。"
                elif dg:
                    sub_para = f"完成数据提取后，服务端按照统一格式完成数据组装封装，将{dg}原路返回至客户端，供前端进行页面渲染与内容展示。"
                else:
                    sub_para = f"完成数据提取后，服务端按照统一格式完成数据组装封装，将处理结果原路返回至客户端。"
            else:
                # Process-type
                if dg and attr_text:
                    sub_para = f"服务端按照业务规则完成{dg}的处理与状态更新，处理内容包括{attr_text}。"
                elif dg:
                    sub_para = f"服务端按照业务规则完成{dg}的处理与状态更新。"
                else:
                    sub_para = f"服务端按照业务规则完成{desc}。"

            # Prepend number
            if i < len(num):
                sub_para = f"{num[i]} {desc}\n{sub_para}"
            else:
                sub_para = f"  {desc}\n{sub_para}"

            paragraphs.append(sub_para)

        return paragraphs

    # Build fp_prd
    fp_prd = {}
    for fp_name, entry in fp_data.items():
        if not entry['subs']:
            continue
        fp_prd[fp_name] = gen_paragraphs(fp_name, entry)

    print(f"[H5 Rewrite] Built PRD paragraphs for {len(fp_prd)} function processes")

    # ============================================================
    # Phase 3: Replace H5 content in Word
    # ============================================================

    print("[H5 Rewrite] Loading Word document...")
    doc = Document(input_docx)

    def make_prd_paragraph(text, first_line_indent=True, bold=False):
        """Create a styled paragraph element"""
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')

        # Indent
        ind = OxmlElement('w:ind')
        if first_line_indent:
            ind.set(qn('w:firstLineChars'),'200')
            ind.set(qn('w:firstLine'),'420')
        else:
            ind.set(qn('w:firstLineChars'),'0')
            ind.set(qn('w:firstLine'),'0')
        pPr.append(ind)

        # Spacing
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:before'),'60')
        sp.set(qn('w:after'),'60')
        pPr.append(sp)

        p.append(pPr)

        # Run
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        for a in [qn('w:eastAsia'),qn('w:ascii'),qn('w:hAnsi')]:
            rFonts.set(a, "宋体")
        rPr.append(rFonts)
        for tag in ['w:sz','w:szCs']:
            el = OxmlElement(tag)
            el.set(qn('w:val'),'21')
            rPr.append(el)
        if bold:
            b_el = OxmlElement('w:b')
            rPr.append(b_el)
        r.append(rPr)

        t = OxmlElement('w:t')
        t.set('{http://www.w3.org/XML/1998/namespace}space','preserve')
        t.text = text
        r.append(t)
        p.append(r)

        return p

    # Collect H5 entries
    h5_entries = [(i, p.text.strip()) for i, p in enumerate(doc.paragraphs)
                  if p.style.name == 'Heading 5' and p.text.strip() in fp_prd]
    print(f"[H5 Rewrite] H5 entries to rewrite: {len(h5_entries)}")

    # Process from back to front
    rewritten = 0
    for idx in range(len(h5_entries)-1, -1, -1):
        h5_i, fp_name = h5_entries[idx]
        prd_paragraphs = fp_prd[fp_name]

        # Refresh paragraph list after modifications
        paras = doc.paragraphs
        if h5_i >= len(paras):
            continue
        h5_elem = paras[h5_i]._element

        # ---- Delete ALL content between this H5 and next heading ----
        parent = h5_elem.getparent()
        to_remove = []
        next_sib = h5_elem.getnext()

        while next_sib is not None:
            # Check if this sibling is a heading → stop
            if next_sib.tag == qn('w:p'):
                pPr = next_sib.find(qn('w:pPr'))
                if pPr is not None:
                    pStyle = pPr.find(qn('w:pStyle'))
                    if pStyle is not None:
                        val = pStyle.get(qn('w:val'))
                        if val and val.startswith('Heading'):
                            break
            # Everything else (paragraphs, tables, etc.) → remove
            to_remove.append(next_sib)
            next_sib = next_sib.getnext()

        for elem in to_remove:
            parent.remove(elem)

        # ---- Insert new paragraphs after H5 ----
        last_elem = h5_elem
        for i, text in enumerate(prd_paragraphs):
            p_elem = make_prd_paragraph(text, first_line_indent=True, bold=False)
            last_elem.addnext(p_elem)
            last_elem = p_elem

        rewritten += 1
        if rewritten % 500 == 0:
            print(f"  [H5 Rewrite] {rewritten}/{len(h5_entries)}...")

    doc.save(output_docx)
    print(f"\n[H5 Rewrite] Complete! {rewritten}/{len(h5_entries)} function processes rewritten.")
    print(f"[H5 Rewrite] Final doc: {output_docx}")
